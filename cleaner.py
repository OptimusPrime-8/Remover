import os
import re
from docx import Document

# Folder containing the Word file(s)
folder_path = r"C:\Users\Madhan\Write Final"

# Expanded Italian accented character replacement map
char_map = {
    "à": "a", "À": "A",
    "è": "e", "È": "E",
    "é": "e", "É": "E",
    "ì": "i", "Ì": "I",
    "ò": "o", "Ò": "O",
    "ù": "u", "Ù": "U",
    "—":", "
}

# Words that should remain untouched
exceptional_words = {"μ", "µ"}  # Add more as needed

def clean_text(text, location="", table_mode=False):
    # Skip if text is exactly an exceptional word
    if text.strip() in exceptional_words:
        return text
    
    changes = []
    
    # Replace non-breaking space
    if "\u00A0" in text:
        changes.append("non-breaking space")
        text = text.replace("\u00A0", " ")
    
    # Replace Italian accented characters
    for k, v in char_map.items():
        if k in text:
            changes.append(k)
            text = text.replace(k, v)
    
    # Remove all other non-ASCII except bullet (•) and comma (,)
    text, removed_chars = remove_unwanted_specials(text)
    if removed_chars:
        changes.extend(removed_chars)
    
    # Print changes
    if changes:
        print(f"{location}: {', '.join(sorted(set(changes)))}")
    
    return text

def remove_unwanted_specials(text):
    removed_chars = []
    new_text = ""
    for ch in text:
        # Keep ASCII chars, bullet • (U+2022), and comma
        if ch.isascii() or ch == "•" or ch == ",":
            new_text += ch
        else:
            removed_chars.append(ch)
    return new_text, removed_chars

# Process all .docx files
for file_name in os.listdir(folder_path):
    if file_name.lower().endswith(".docx"):
        file_path = os.path.join(folder_path, file_name)
        print(f"Processing file: {file_name}")
        
        doc = Document(file_path)
        
        # Clean paragraphs
        for i, para in enumerate(doc.paragraphs):
            para.text = clean_text(para.text, f"Paragraph {i+1}")
        
        # Clean tables
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    cell.text = clean_text(cell.text, f"Table {ti+1} Row {ri+1} Col {ci+1}", table_mode=True)
        
        doc.save(file_path)
        print(f"Overwritten: {file_path}\n{'-'*50}")

print("All files processed.")
